from typing import List, Dict
import logging
from src.rag_app.core.interfaces.chunk_strategy_interface import ChunkStrategyInterface
from src.rag_app.core.interfaces.document_interface import Chunk
from docx import Document
import json
import os
from pypdf import PdfReader
import re

logger = logging.getLogger(__name__)

class StructuredDocumentStrategy(ChunkStrategyInterface):
    def __init__(self, chunk_size: int = 1000, overlap: int = 100, max_chunk_size: int = 4000, min_chunk_size: int = 350):
        self._strategy_name = "Structured Document"
        self.chunk_size = chunk_size
        self.max_chunk_size = max_chunk_size
        self.overlap = overlap if overlap is not None else 0
        self.min_chunk_size = min_chunk_size

    @property
    def strategy_name(self) -> str:
        return self._strategy_name

    def get_parameters(self) -> Dict[str, any]:
        return {
            "max_chunk_size": self.max_chunk_size,
            "min_chunk_size": self.min_chunk_size,
            "chunk_size": self.chunk_size,
            "overlap": self.overlap
        }

    def _split_content_with_overlap(self, title: str, content: List[str]) -> List[str]:
        """
        Split content into chunks respecting chunk_size and overlap.
        
        Args:
            title: Document title to be included with the first chunk
            content: List of content strings to be chunked
            
        Returns:
            List of text chunks with proper overlap and minimum size constraints
        """
        # Prepare full text with title
        full_text = f"{title}\n\n" + "\n".join(content)
        chunks = []
        start = 0
        
        # Ensure overlap is an integer and not None
        overlap = max(0, self.overlap if self.overlap is not None else 0)
        
        while start < len(full_text):
            # Get chunk of specified size
            chunk = full_text[start:start + self.chunk_size]
            
            # Ensure chunk is not empty
            if not chunk:
                break
                
            chunks.append(chunk)
            # Move start position by (chunk_size - overlap), ensuring at least 1 character advance
            start += max(self.chunk_size - overlap, 1)
            
        # Handle last chunk if it's too small
        if len(chunks) >= 2 and len(chunks[-1]) < self.min_chunk_size:
            # Merge the last chunk with the previous one
            merged_chunk = chunks[-2] + chunks[-1]
            chunks = chunks[:-2] + [merged_chunk]
            
        return chunks

    def _is_content_relevant(self, content: str) -> bool:
        """
        Check if the content is relevant (not empty, not only spaces or punctuation).
        
        Args:
            content: String content to validate
            
        Returns:
            bool: True if content is relevant, False otherwise
        """
        # Remove spaces and common punctuation
        cleaned_content = ''.join(char for char in content if char.isalnum())
        return len(cleaned_content) > 0

    def chunk_text(self, content: str, document_id: str, doc_path: str) -> List[Chunk]:
        logger.info("Starting chunk_text process for document_id: %s", document_id)
        
        if doc_path is None:
            logger.error("doc_path must be provided for StructuredDocumentChunker")
            return []

        logger.info("Extracting document structure from path: %s", doc_path)
        doc_structure = self.extract_docx_with_structure(doc_path)
        chunks = []
        chunk_id = 0

        # Process default section if it exists and handle short content
        if "default" in doc_structure:
            logger.info("Processing default section")
            default_content = "\n".join(doc_structure["default"]["content"])
            
            if len(default_content) < self.min_chunk_size and doc_structure["sections"]:
                logger.info("Default section is too short (%d chars), merging with first section", 
                           len(default_content))
                # Merge with first non-skipped section
                for section in doc_structure["sections"]:
                    if not section.get("skip", False):
                        section["content"] = doc_structure["default"]["content"] + section["content"]
                        section["tables"] = doc_structure["default"]["tables"] + section["tables"]
                        section["images"] = doc_structure["default"]["images"] + section["images"]
                        break
            else:
                chunks.append(self._create_chunk(
                    content=doc_structure["default"]["content"],
                    document_id=document_id,
                    chunk_id=chunk_id,
                    breadcrumb="root",
                    heading="default",
                    parents=[],
                    tables=doc_structure["default"]["tables"],
                    images=doc_structure["default"]["images"]
                ))
                chunk_id += 1

        # Create a mapping of sections to their parents
        logger.info("Building section hierarchy")
        section_parents = self._build_section_hierarchy(doc_structure["sections"])
        logger.info(f"Printing section parents {section_parents}")

        previous_section = None
        merged_content = []
        i = 0
        
        # First pass: identify and merge small sections
        while i < len(doc_structure["sections"]):
            section = doc_structure["sections"][i]
            
            if not any(self._is_content_relevant(text) for text in section["content"]):
                logger.info("Skipping section with no relevant content: %s", section["title"])
                section["skip"] = True
                i += 1
                continue
                
            current_content = "\n".join(section["content"])
            logger.info("Section %s", section["title"])
            logger.info("Content %s", current_content)
            logger.info("Content len %s", len(current_content))

            if len(current_content) < self.min_chunk_size and not section.get("skip", False):
                logger.info("Section %s is smaller than min_chunk_size, attempting to merge", section["title"])
                
                merged_section = None
                sections_to_merge = 1
                
                # Keep trying to merge with subsequent sections until we reach min_chunk_size
                while (len(current_content) < self.min_chunk_size and 
                       i + sections_to_merge < len(doc_structure["sections"])):
                    next_section = doc_structure["sections"][i + sections_to_merge]
                    
                    # Skip already merged sections
                    if next_section.get("skip", False):
                        sections_to_merge += 1
                        continue
                        
                    # Create merged section
                    merged_section = {
                        "title": " + ".join([section["title"]] + 
                                          [doc_structure["sections"][j]["title"] 
                                           for j in range(i + 1, i + sections_to_merge + 1)
                                           if not doc_structure["sections"][j].get("skip", False)]),
                        "level": min(section["level"], 
                                   min(s["level"] for s in doc_structure["sections"][i + 1:i + sections_to_merge + 1]
                                       if not s.get("skip", False))),
                        "content": [],
                        "tables": [],
                        "images": []
                    }
                    
                    # Accumulate content and assets from all sections being merged
                    for j in range(i, i + sections_to_merge + 1):
                        if not doc_structure["sections"][j].get("skip", False):
                            merged_section["content"].extend(doc_structure["sections"][j]["content"])
                            merged_section["tables"].extend(doc_structure["sections"][j]["tables"])
                            merged_section["images"].extend(doc_structure["sections"][j]["images"])
                    
                    current_content = "\n".join(merged_section["content"])
                    sections_to_merge += 1
                    
                    # Break if we've reached the end of the document
                    if i + sections_to_merge >= len(doc_structure["sections"]):
                        break
                
                if merged_section:
                    # Mark sections as skipped
                    for j in range(i + 1, min(i + sections_to_merge, len(doc_structure["sections"]))):
                        doc_structure["sections"][j]["skip"] = True
                    
                    merged_content.append((i, merged_section))
                    i += sections_to_merge
                    continue
                
                # If we couldn't merge forward and we have a previous section, try merging backward
                elif previous_section is not None and not previous_section.get("skip", False):
                    merged_section = {
                        "title": f"{previous_section['title']} + {section['title']}",
                        "level": min(previous_section["level"], section["level"]),
                        "content": previous_section["content"] + section["content"],
                        "tables": previous_section["tables"] + section["tables"],
                        "images": previous_section["images"] + section["images"]
                    }
                    # Replace the previous section with merged content
                    merged_content.append((i-1, merged_section))
            
            previous_section = section
            i += 1

        # Apply merges to the original sections list
        for index, merged_section in reversed(merged_content):
            doc_structure["sections"][index] = merged_section

        # Second pass: process sections and create chunks
        for section in doc_structure["sections"]:
            if section.get("skip", False):
                logger.info("Skipping previously marked section: %s", section["title"])
                continue
                
            logger.info("Processing section: %s", section["title"])
            breadcrumb = self._create_breadcrumb(section, section_parents)
            
            # Split section content into overlapping chunks
            logger.info("Splitting content for section: %s", section["title"])
            content_chunks = self._split_content_with_overlap(section["title"], section["content"])
            
            for idx, content_chunk in enumerate(content_chunks):
                if not self._is_content_relevant(content_chunk):
                    logger.info("Skipping irrelevant chunk for section: %s", section["title"])
                    continue
                    
                logger.info("Creating chunk %d for section: %s", idx + 1, section["title"])
                chunk = self._create_chunk(
                    content=[content_chunk],
                    document_id=document_id,
                    chunk_id=chunk_id,
                    breadcrumb=f"{breadcrumb} (part {idx + 1}/{len(content_chunks)})",
                    heading=section["title"],
                    parents=section_parents.get(section["title"], []),
                    tables=section["tables"],
                    images=section["images"]
                )
                chunks.append(chunk)
                chunk_id += 1

        # Add chunk analysis logging before returning
        logger.info("Chunk analysis for document: %s", document_id)
        
        # Calculate average length
        chunk_lengths = [len(chunk.content) for chunk in chunks]
        avg_length = sum(chunk_lengths) / len(chunks) if chunks else 0
        logger.info("Average chunk length: %d characters", avg_length)
        
        # Sort chunks by length for top/bottom analysis
        chunks_with_lengths = [(chunk, len(chunk.content)) for chunk in chunks]
        sorted_chunks = sorted(chunks_with_lengths, key=lambda x: x[1], reverse=True)
        
        # Log top 5 chunks
        logger.info("Top 5 longest chunks:")
        for chunk, length in sorted_chunks[:5]:
            logger.info("- Length: %d, Heading: %s", length, chunk.metadata['heading'])
            logger.info("  Content preview: %s...", chunk.content[:100])
            
        # Log bottom 5 chunks
        logger.info("Bottom 5 shortest chunks:")
        for chunk, length in sorted_chunks[-5:]:
            logger.info("- Length: %d, Heading: %s", length, chunk.metadata['heading'])
            logger.info("  Content preview: %s...", chunk.content[:100])

        logger.info("Completed chunk_text process for document_id: %s", document_id)
        return chunks

    def extract_docx_with_structure(self, file_path: str) -> Dict:
        """Extract structure from a document file (DOCX or PDF)."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension in ['.doc', '.docx']:
            return self._extract_docx_structure(file_path)
        elif file_extension == '.pdf':
            return self._extract_pdf_structure(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_extension}")
            return {"file": file_path, "sections": []}

    def _extract_docx_structure(self, docx_path: str) -> Dict:
        """Extract structure specifically from DOCX files."""
        doc = Document(docx_path)
        structure = {"file": docx_path, "sections": []}
        current_section = None

        def get_default_section():
            if "default" not in structure:
                structure["default"] = {"content": [], "tables": [], "images": []}
            return structure["default"]

        body = doc.element.body
        current_table_index = 0

        for child in body.iter():
            if child.tag.endswith('}p'):
                for paragraph in doc.paragraphs:
                    if paragraph._element is child:
                        text = paragraph.text.strip()
                        if not text:
                            continue

                        style_name = paragraph.style.name

                        if style_name.startswith("Heading"):
                            heading_level = int(style_name.split()[-1])
                            current_section = {
                                "title": text,
                                "level": heading_level,
                                "content": [],
                                "tables": [],
                                "images": []
                            }
                            structure["sections"].append(current_section)
                        else:
                            target_section = current_section if current_section else get_default_section()
                            target_section["content"].append(text)
                        break

            elif child.tag.endswith('}tbl'):
                if current_table_index < len(doc.tables):
                    table = doc.tables[current_table_index]
                    table_content = []
                    for row in table.rows:
                        row_content = [cell.text.strip() for cell in row.cells]
                        table_content.append(row_content)

                    target_section = current_section if current_section else get_default_section()
                    target_section["content"].append(f"[TABLE_{current_table_index}]")
                    target_section["tables"].append({"id": current_table_index, "content": table_content})
                    current_table_index += 1

            elif child.tag.endswith('}drawing'):
                blip = child.find('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                if blip is not None:
                    rel_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rel_id and rel_id in doc.part.rels:
                        image_path = doc.part.rels[rel_id].target_ref
                        target_section = current_section if current_section else get_default_section()
                        target_section["content"].append(f"![Image]({image_path})")
                        target_section["images"].append({"id": image_path, "content": image_path})

        return structure

    def _extract_pdf_structure(self, pdf_path: str) -> Dict:
        """Extract structure from PDF files using heading detection and content organization."""
        structure = {"file": pdf_path, "sections": []}
        current_section = None
        
        try:
            reader = PdfReader(pdf_path)
            
            # Initialize default section
            structure["default"] = {"content": [], "tables": [], "images": []}
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                
                # Split text into lines
                lines = text.split('\n')
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Heading detection heuristics
                    is_heading = self._is_potential_heading(line)
                    
                    if is_heading:
                        # Determine heading level based on characteristics
                        heading_level = self._determine_heading_level(line)
                        
                        current_section = {
                            "title": line,
                            "level": heading_level,
                            "content": [],
                            "tables": [],
                            "images": []
                        }
                        structure["sections"].append(current_section)
                    else:
                        # Add content to current section or default
                        target_section = current_section if current_section else structure["default"]
                        target_section["content"].append(line)
                
                # Extract images if available
                if page.images:
                    target_section = current_section if current_section else structure["default"]
                    for i, image in enumerate(page.images):
                        image_id = f"page_{page_num}_image_{i}"
                        target_section["content"].append(f"![Image]({image_id})")
                        target_section["images"].append({"id": image_id, "content": image_id})
            
            return structure
            
        except Exception as e:
            logger.error(f"Error processing PDF file: {str(e)}")
            return {"file": pdf_path, "sections": []}

    def _is_potential_heading(self, line: str) -> bool:
        """
        Determine if a line is likely to be a heading based on various heuristics.
        """
        # Heading characteristics
        characteristics = [
            len(line) < 200,  # Headings are usually shorter
            line.strip().endswith(('.', ':', '?')),  # Common heading punctuation
            any(char.isupper() for char in line),  # Contains uppercase letters
            not line.startswith('    '),  # Not heavily indented
            bool(re.match(r'^[0-9.]+\s+', line))  # Numbered sections
        ]
        
        # If line meets multiple characteristics, it's likely a heading
        return sum(characteristics) >= 3

    def _determine_heading_level(self, line: str) -> int:
        """
        Determine the heading level based on text characteristics.
        """
        # Default level
        level = 1
        
        # Check for numbered sections (e.g., "1.2.3")
        if re.match(r'^[0-9.]+\s+', line):
            level = line.split()[0].count('.') + 1
        
        # Check indentation
        indent = len(line) - len(line.lstrip())
        level += indent // 2
        
        # Ensure level is within reasonable bounds
        return max(1, min(6, level))

    def _build_section_hierarchy(self, sections: List[Dict]) -> Dict[str, List[str]]:
        section_parents = {}
        section_stack = []

        for section in sections:
            while section_stack and section_stack[-1]["level"] >= section["level"]:
                section_stack.pop()

            section_parents[section["title"]] = [s["title"] for s in reversed(section_stack)]
            section_stack.append(section)

        return section_parents

    def _create_breadcrumb(self, section: Dict, section_parents: Dict[str, List[str]]) -> str:
        parents = section_parents.get(section["title"], [])
        reversed_parents = list(reversed(parents))
        return " > ".join(reversed_parents + [section["title"]])

    def _create_chunk(self, content: List[str], document_id: str, 
                      chunk_id: int, breadcrumb: str, heading: str, 
                      parents: List[str], tables: List[Dict], images: List[Dict]) -> Chunk:
        full_content = "\n".join(content)

        return Chunk(
            document_id=document_id,
            chunk_id=f"{document_id}_chunk_{chunk_id}",
            metadata={
                "breadcrumb": breadcrumb,
                "chunk_id": f"{document_id}_chunk_{chunk_id}",
                "heading": heading,
                "parents": " > ".join(parents)#,
                #"tables": tables,
                #"images": images
            },
            content=full_content
        ) 

    async def format_result(self, data_path, combined_results: List[dict], result_domains: List[str]) -> List[dict]:
        """Format and enhance results based on chunk strategy."""

        # 1. Load chunks with validation
        base_chunks_dir = os.path.join(data_path, '../chunks')
        all_chunks = []
        
        if not os.path.exists(base_chunks_dir):
            logger.error(f"Chunks directory not found: {base_chunks_dir}")
            return combined_results

        structured_dirs = [d for d in os.listdir(base_chunks_dir) 
                        if os.path.isdir(os.path.join(base_chunks_dir, d)) 
                        and d.endswith('_Structured Document')]

        # Load chunks with content validation
        for dir_name in structured_dirs:
            dir_path = os.path.join(base_chunks_dir, dir_name)
            json_files = [f for f in os.listdir(dir_path) if f.endswith('.json')]
            
            for json_file in json_files:
                try:
                    with open(os.path.join(dir_path, json_file), 'r', encoding='utf-8') as f:
                        chunks = json.load(f)
                        # Validate chunk structure
                        for chunk in chunks:
                            if not all(k in chunk for k in ['chunk_id', 'content', 'metadata']):
                                logger.warning(f"Invalid chunk structure in {json_file}: {chunk.keys()}")
                                continue
                            all_chunks.append(chunk)
                except Exception as e:
                    logger.error(f"Error loading {json_file}: {str(e)}")

        # Create chunks dictionary with validation
        chunks_dict = {}
        for chunk in all_chunks:
            chunk_id = chunk['chunk_id']
            if chunk_id in chunks_dict:
                logger.warning(f"Duplicate chunk_id found: {chunk_id}")
            chunks_dict[chunk_id] = chunk

        # Add before processing results:
        processed_breadcrumbs = set()  # Track which breadcrumbs we've already processed
        formatted_results = []

        for idx, (result, domain) in enumerate(zip(combined_results, result_domains)):
            try:
                
                metadata = result.get('metadata', {})
                if isinstance(metadata, list) and metadata:
                    metadata = metadata[0]
                
                chunk_id = metadata.get('chunk_id', '')
                if not chunk_id or chunk_id not in chunks_dict:
                    logger.warning(f"Chunk ID {chunk_id} not found in chunks_dict")
                    formatted_results.append(result)
                    continue

                # Get current chunk content from the correct source
                current_chunk = chunks_dict[chunk_id]
                #logger.info(f"Current chunk {current_chunk['metadata']['chunk_id']}")

                breadcrumb = current_chunk['metadata']['breadcrumb']
                if '(part ' not in breadcrumb:
                    logger.warning(f"No part information found in breadcrumb: {breadcrumb}")
                    formatted_results.append(result)
                    continue
                
                # Extract base breadcrumb (without part information)
                base_breadcrumb = breadcrumb.split(' (part ')[0] + current_chunk['metadata']['heading']
                # Skip if we've already processed this breadcrumb
                if base_breadcrumb in processed_breadcrumbs:
                    logger.debug(f"Skipping duplicate breadcrumb: {base_breadcrumb}")
                    continue
                
                processed_breadcrumbs.add(base_breadcrumb)

                part_info = breadcrumb.split('(part ')[-1].split(')')[0]
                current_part, total_parts = map(int, part_info.split('/'))
                document_id = current_chunk['metadata']['document_id']
                #logger.info(f"Current part {current_part}, total parts {total_parts}")
                
                # Create section header
                breadcrumb_without_part = breadcrumb.split(' (part ')[0]
                section_header = f"\n\n\Ενότητα: {breadcrumb_without_part} > {current_chunk['metadata']['heading']}"
                
                contents = []
                is_first_chunk = (current_part == 1)
                
                document_id = current_chunk['metadata']['document_id']
                # Extract the base and number from document_id
                base_id, current_num = current_chunk["chunk_id"].rsplit('_', 1)
                current_num = int(current_num)

                #logger.info(f"base_id {base_id}")
                #logger.info(f"current_num {current_num}")
                max_parts_threshold = 99
                contents = []

                if total_parts <= max_parts_threshold:

                    first_chunk = current_num - current_part + 1
                    # Get all parts in sequence
                    for i in range(total_parts):
                        logging.debug(f"Appending {base_id}_{first_chunk + i}")
                        #fetch_num = parts_before + i
                        fetch_id = f"{base_id}_{first_chunk + i}"
                        contents.append(chunks_dict[fetch_id]['content'])

                # Combine contents with section header and continuation marker if needed
                combined_content = section_header + "\n"

                combined_content += '\n'.join(contents)

                formatted_result = {
                    'document_id': current_chunk['metadata']['document_id'],
                    'chunk_id': chunk_id,
                    'metadata': current_chunk['metadata'],
                    'content': combined_content,
                    'document': combined_content,
                    'distance': result.get('distance'),
                    'domain': domain
                }
                
                formatted_results.append(formatted_result)
                logger.info(f"Successfully processed chunk {chunk_id} with distance {result.get('distance')}")
                

            except Exception as e:
                logger.error(f"Error processing result: {str(e)}", exc_info=True)
                formatted_results.append(result)

        #logger.info(f"Final formatted results length: {len(formatted_results)}")
        return formatted_results